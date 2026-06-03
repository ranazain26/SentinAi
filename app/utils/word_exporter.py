import docx
from docx.shared import Pt, RGBColor
from datetime import datetime

def export_word(analysis: dict, output_path: str = "sentinai_report.docx"):
    """Generate a native Word document report from the analysis dictionary."""
    doc = docx.Document()
    
    # Title & Meta
    title = doc.add_heading('SentinAI Security Audit Report', 0)
    
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    doc.add_paragraph(f"File: {analysis.get('filename', 'unknown')}")
    doc.add_paragraph(f"Generated: {now}")
    doc.add_paragraph(f"Language: {analysis.get('parse', {}).get('language', 'unknown').capitalize()}")
    
    # ── Code Health Score Table ──
    doc.add_heading('Code Health Score', level=1)
    score_data = analysis.get("score", {})
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Value'
    
    metrics = [
        ("Score", f"{score_data.get('score', 0)}/100"),
        ("Grade", str(score_data.get('grade', 'N/A'))),
        ("Status", str(score_data.get('status', 'Unknown'))),
        ("Lines of Code", str(score_data.get('lines_of_code', 0))),
        ("Critical Issues", str(score_data.get('critical_issues', 0))),
        ("High Issues", str(score_data.get('high_issues', 0))),
        ("Total Vulnerabilities", str(score_data.get('total_vulnerabilities', 0))),
        ("Total Bugs", str(score_data.get('total_bugs', 0))),
        ("Secrets Detected", str(score_data.get('total_secrets', 0)))
    ]
    
    for metric, value in metrics:
        row_cells = table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = value

    # ── Security Vulnerabilities ──
    security = analysis.get("security", [])
    doc.add_heading(f'Security Vulnerabilities ({len(security)})', level=1)
    
    if not security:
        doc.add_paragraph("No security vulnerabilities detected.")
    else:
        rag_refs = analysis.get("rag_references", {})
        for i, v in enumerate(security, 1):
            doc.add_heading(f"{i}. {v['vulnerability']} — {v['severity']}", level=2)
            p = doc.add_paragraph()
            p.add_run("Line: ").bold = True
            p.add_run(f"{v['line']}\n")
            p.add_run("CWE: ").bold = True
            p.add_run(f"{v.get('cwe', 'N/A')}\n")
            p.add_run("Code: ").bold = True
            p.add_run(f"{v['snippet']}\n")
            p.add_run("Description: ").bold = True
            p.add_run(f"{v['description']}")
            
            refs = rag_refs.get(v["vulnerability"], [])
            if refs:
                p.add_run("\nReferences:\n").bold = True
                for r in refs:
                    p.add_run(f"- [{r.get('id', 'N/A')}] {r.get('name', '')}: {r.get('remediation', '')}\n")

    # ── Bugs & Secrets ──
    bugs = analysis.get("bugs", [])
    doc.add_heading(f'Bugs & Code Quality Issues ({len(bugs)})', level=1)
    if not bugs:
        doc.add_paragraph("No bugs detected.")
    else:
        for i, b in enumerate(bugs, 1):
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{b['bug']} [{b['severity']}] ").bold = True
            p.add_run(f"(Line {b['line']}): {b['description']}")

    ner = analysis.get("ner", [])
    doc.add_heading(f'Sensitive Entities ({len(ner)})', level=1)
    if not ner:
        doc.add_paragraph("No sensitive entities detected.")
    else:
        for entity in ner:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{entity['type']} [{entity['severity']}] ").bold = True
            p.add_run(f"(Line {entity['line']})")

    # ── Auto-Generated Documentation ──
    docs = analysis.get("documentation", {})
    doc.add_heading('Auto-Generated Documentation', level=1)
    doc.add_heading('Summary', level=2)
    doc.add_paragraph(docs.get("summary", "No summary available."))
    
    doc.add_heading('Purpose', level=2)
    doc.add_paragraph(docs.get("purpose", "N/A"))
    
    doc.save(output_path)
    return output_path